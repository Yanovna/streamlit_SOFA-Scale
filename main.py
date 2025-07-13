import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime
from scales.sofa import calculate_sofa_scores

def init_session_state():
    if 'scores' not in st.session_state:
        st.session_state.scores = {
            'respiratory': 0,
            'nervous': 0,
            'cardiovascular': 0,
            'liver': 0,
            'coagulation': 0,
            'renal': 0
        }
    if 'total_score' not in st.session_state:
        st.session_state.total_score = 0
    if 'sofa_result' not in st.session_state:
        st.session_state.sofa_result = None

def generate_report():
    doc = Document()
    doc.add_heading('Отчёт по шкале SOFA', level=1)
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Общий балл SOFA: {st.session_state.total_score}")
    doc.add_heading("Детализация баллов:", level=2)

    def format_score(score):
        if score == 1:
            return f"{score} балл"
        elif 2 <= score <= 4:
            return f"{score} балла"
        else:
            return f"{score} баллов"

    doc.add_paragraph(f"1. Респираторный индекс: {format_score(st.session_state.scores['respiratory'])}")
    doc.add_paragraph(f"2. Шкала комы Глазго: {format_score(st.session_state.scores['nervous'])}")
    doc.add_paragraph(f"3. СCC: {format_score(st.session_state.scores['cardiovascular'])}")
    doc.add_paragraph(f"4. Билирубин: {format_score(st.session_state.scores['liver'])}")
    doc.add_paragraph(f"5. Тромбоциты: {format_score(st.session_state.scores['coagulation'])}")
    doc.add_paragraph(f"6. Почечная функция: {format_score(st.session_state.scores['renal'])}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def show_results():
    report = generate_report()
    st.download_button(
        label="📥 Скачать отчёт (DOCX)",
        data=report,
        file_name=f"SOFA_отчёт_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def main():
    st.title("Шкала SOFA (Sequential Organ Failure Assessment)")
    init_session_state()

    with st.form("sofa_form"):
        params = {
            'pao2': st.slider('PaO2 (мм рт.ст.)', 1, 250, 100),
            'fio2': st.slider('FiO2 (%)', 21, 100, 21),
            'gcs': st.slider('Шкала Глазго', 3, 15, 15),
            'cardiovascular': st.selectbox('Сердечно-сосудистая система',
                                       ['Норма', 'АДср < 70', 'Допамин ≤ 5',
                                        'Допамин > 5', 'Допамин > 15']),
            'bilirubin': st.slider('Билирубин (мкмоль/л)', 0, 400, 0),
            'platelets': st.slider('Тромбоциты (×10⁹/л)', 0, 500, 0),
            'renal': st.selectbox('Почечная функция',
                              ['< 110 мкмоль/л', '110-170 мкмоль/л',
                               '171-299 мкмоль/л', '300-440 мкмоль/л',
                               '> 440 мкмоль/л'])
        }

        if st.form_submit_button("Рассчитать общий балл SOFA"):
            st.session_state.sofa_result = calculate_sofa_scores(params)

            if st.session_state.sofa_result:
                st.session_state.scores = st.session_state.sofa_result['scores']
                st.session_state.total_score = st.session_state.sofa_result['total']

    if st.session_state.sofa_result:
        result = st.session_state.sofa_result
        st.subheader(f"Общий балл: {result['total']}")
        st.write(f"Диагноз: {result['diagnosis']}")
        show_results()

        if st.button("Сбросить все данные"):
            init_session_state()
            st.rerun()

if __name__ == "__main__":
    main()